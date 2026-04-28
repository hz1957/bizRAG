#!/usr/bin/env python3
from __future__ import annotations

import argparse
from collections import defaultdict
import hashlib
import json
import logging
import os
import random
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from contextlib import contextmanager
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable, List, Sequence
from xml.etree import ElementTree as ET

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

DEFAULT_WORKSPACE_DIR = PROJECT_ROOT / "rag_eval_workspace"
DOTENV_PATH = PROJECT_ROOT / ".env"

import pandas as pd
from dotenv import load_dotenv
from huggingface_hub import InferenceClient
from openai import OpenAI

try:
    from bizrag.servers.corpus.corpus import (
        _read_docx_text,
        _read_via_office_convert,
        clean_text,
        reflow_paragraphs,
        suppress_stdout,
    )
except ModuleNotFoundError:
    @contextmanager
    def suppress_stdout() -> Iterable[None]:
        stdout_fd = sys.stdout.fileno()
        saved_stdout_fd = os.dup(stdout_fd)
        try:
            devnull = os.open(os.devnull, os.O_WRONLY)
            os.dup2(devnull, stdout_fd)
            os.close(devnull)
            yield
        finally:
            os.dup2(saved_stdout_fd, stdout_fd)
            os.close(saved_stdout_fd)


    def _local_name(tag: str) -> str:
        return tag.split("}", 1)[1] if "}" in tag else tag


    def _read_docx_text_zip(fp: str) -> str | None:
        try:
            with zipfile.ZipFile(fp) as zf:
                if "word/document.xml" not in zf.namelist():
                    return None
                xml_bytes = zf.read("word/document.xml")
        except (zipfile.BadZipFile, OSError):
            return None

        try:
            root = ET.fromstring(xml_bytes)
        except ET.ParseError:
            return None

        paragraphs: list[str] = []
        for para in root.iter():
            if _local_name(para.tag) != "p":
                continue
            buf: list[str] = []
            for node in para.iter():
                lname = _local_name(node.tag)
                if lname == "t" and node.text:
                    buf.append(node.text)
                elif lname == "tab":
                    buf.append("\t")
                elif lname in {"br", "cr"}:
                    buf.append("\n")
            para_text = "".join(buf).strip()
            if para_text:
                paragraphs.append(para_text)
        return "\n".join(paragraphs)


    def _read_docx_text(fp: str) -> str | None:
        try:
            from docx import Document
        except ImportError:
            return _read_docx_text_zip(fp)

        try:
            doc = Document(fp)
            full_text = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            text = "\n".join(full_text).strip()
            return text or _read_docx_text_zip(fp)
        except Exception:
            return _read_docx_text_zip(fp)


    def _find_office_cmd() -> str | None:
        return shutil.which("soffice") or shutil.which("libreoffice")


    def _convert_to_docx_with_office(fp: str, out_dir: str, office_cmd: str) -> str | None:
        cmd = [office_cmd, "--headless", "--convert-to", "docx", "--outdir", out_dir, fp]
        try:
            subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
                timeout=60,
            )
        except Exception:
            return None

        expected = Path(out_dir) / f"{Path(fp).stem}.docx"
        if expected.exists():
            return str(expected)
        for candidate in Path(out_dir).glob("*.docx"):
            return str(candidate)
        return None


    def _read_with_text_extractor(fp: str, command_name: str) -> str | None:
        extractor_cmd = shutil.which(command_name)
        if not extractor_cmd:
            return None
        try:
            proc = subprocess.run(
                [extractor_cmd, fp],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
                timeout=60,
            )
        except Exception:
            return None
        text = proc.stdout.decode("utf-8", errors="ignore").strip()
        return text or None


    def _convert_to_docx_with_textutil(fp: str, out_dir: str) -> str | None:
        textutil_cmd = shutil.which("textutil")
        if not textutil_cmd:
            return None
        output_path = Path(out_dir) / f"{Path(fp).stem}.docx"
        try:
            subprocess.run(
                [textutil_cmd, "-convert", "docx", "-output", str(output_path), fp],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
                timeout=60,
            )
        except Exception:
            return None
        return str(output_path) if output_path.exists() else None


    def _read_via_office_convert(fp: str) -> str | None:
        with tempfile.TemporaryDirectory(prefix="rag_eval_docx_") as tmpdir:
            office_cmd = _find_office_cmd()
            out_path = _convert_to_docx_with_office(fp, tmpdir, office_cmd) if office_cmd else None
            if not out_path:
                for text_extractor in ("antiword", "catdoc"):
                    text = _read_with_text_extractor(fp, text_extractor)
                    if text:
                        return text
                out_path = _convert_to_docx_with_textutil(fp, tmpdir)
                if not out_path:
                    return None
            return _read_docx_text(out_path)


    def clean_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\u3000", " ")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


    def reflow_paragraphs(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        end_punct_re = re.compile(r"[。！？!?；;…]\s*[”’」』》）】]*\s*$")
        next_start_re = re.compile(r'^[\u4e00-\u9fff0-9a-zA-Z“"‘’《（(【\[「『<]')

        def merge_lines_within_paragraph(paragraph: str) -> str:
            lines = paragraph.split("\n")
            segments: list[str] = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if not segments:
                    segments.append(line)
                    continue
                prev = segments[-1]
                if prev.endswith("-"):
                    segments[-1] = prev[:-1] + line
                elif end_punct_re.search(prev):
                    segments.append(line)
                else:
                    joiner = "" if re.match(r"^[,.;:!?%)]", line) else " "
                    segments[-1] = prev + joiner + line
            return "\n".join(segments)

        paragraphs = [merge_lines_within_paragraph(part) for part in re.split(r"\n\s*\n", text) if part.strip()]
        merged: list[str] = []
        for paragraph in paragraphs:
            if not merged:
                merged.append(paragraph)
                continue
            prev = merged[-1]
            first_line = paragraph.split("\n", 1)[0].strip()
            if not end_punct_re.search(prev) and next_start_re.search(first_line):
                merged[-1] = f"{prev} {paragraph}".strip()
            else:
                merged.append(paragraph)
        return "\n\n".join(merged).strip()


LOGGER = logging.getLogger("build_rag_eval_dataset")

TEXT_EXTS = {".txt", ".md"}
DOCX_EXTS = {".docx"}
WORD_LEGACY_EXTS = {".doc", ".wps"}
PDFLIKE_EXTS = {".pdf", ".xps", ".oxps", ".epub", ".mobi", ".fb2"}
EXCEL_EXTS = {".xlsx", ".xls"}
JSONL_EXTS = {".jsonl"}

DEFAULT_SEPARATORS = ["\n\n", "\n", ".", " ", ""]
DEFAULT_TEXT_COLUMNS = ["text", "contents", "page_content", "content", "body"]
DEFAULT_SOURCE_COLUMNS = ["source", "source_uri", "url", "path", "file_path"]
DEFAULT_TITLE_COLUMNS = ["title", "name", "doc_title"]
DEFAULT_FILE_NAME_COLUMNS = ["file_name", "filename", "name"]
DEFAULT_SOURCE_TYPE_COLUMNS = ["source_type", "doc_type", "type"]

QA_GENERATION_PROMPT = """
Your task is to write one factoid question and one short answer from the context.
The question must be answerable with a specific, concise factual answer from the context.
The question must look like something a user would type into a search engine.
Do not mention "the passage", "the context", or "the document".
Write the question and answer in the same language as the context.

Provide your answer exactly in this format:

Output:::
Factoid question: (your factoid question)
Answer: (your answer)

Context: {context}
Output:::
""".strip()

CORE_FACTOID_PROMPT = """
Your task is to write one factoid question and one short answer from the context.
Requirements:
- The question must be fully answerable from the context.
- The answer must be unique, concise, and factual.
- Prefer questions with short answers such as a name, date, number, criterion, or short phrase.
- Do not mention "the passage", "the context", or "the document".
- Write the question and answer in the same language as the context.

Provide your answer exactly in this format:

Output:::
Factoid question: (your factoid question)
Answer: (your answer)

Context: {context}
Output:::
""".strip()

CROSS_SENTENCE_PROMPT = """
Your task is to write one factoid question and one short answer from the two contexts.
Requirements:
- The question must require combining information from BOTH Context 1 and Context 2.
- The answer must depend on at least one detail that appears only in Context 1 and at least one detail that appears only in Context 2.
- The question must still have a single concise factual answer.
- Do not write a question that can be answered from Context 1 alone or Context 2 alone.
- Do not rely on duplicated or overlapping text that appears in both contexts.
- Do not mention "the passage", "the context", or "the document".
- Write the question and answer in the same language as the context.

Provide your answer exactly in this format:

Output:::
Factoid question: (your factoid question)
Answer: (your answer)

Context: {context}
Output:::
""".strip()

UNANSWERABLE_PROMPT = """
Your task is to write one plausible user question related to the context that CANNOT be answered from the context alone.
Requirements:
- The question must look realistic and relevant to the context.
- The missing information must not be stated anywhere in the context.
- The question should be understandable on its own.
- Do not mention "the passage", "the context", or "the document".
- The answer must be exactly NO_ANSWER.
- Write the question in the same language as the context.

Provide your answer exactly in this format:

Output:::
Factoid question: (your factoid question)
Answer: NO_ANSWER

Context: {context}
Output:::
""".strip()

AMBIGUOUS_REWRITE_PROMPT = """
Your task is to write one factoid question and one short answer from the context.
Requirements:
- The question must be fully answerable from the context.
- The wording can be slightly indirect or colloquial, but the answer must still be unique.
- Do not create a truly multi-answer or ambiguous question.
- Do not mention "the passage", "the context", or "the document".
- Write the question and answer in the same language as the context.

Provide your answer exactly in this format:

Output:::
Factoid question: (your factoid question)
Answer: (your answer)

Context: {context}
Output:::
""".strip()

QUESTION_GROUNDEDNESS_PROMPT = """
You will be given a context and a question.
Score how well the question can be answered unambiguously from the context.
Use a scale of 1 to 5, where 1 means not answerable at all and 5 means clearly answerable.

Provide your answer exactly in this format:

Answer:::
Evaluation: (your rationale)
Total rating: (1 to 5)

Question: {question}
Context: {context}
Answer:::
""".strip()

QUESTION_RELEVANCE_PROMPT = """
You will be given a question.
Score how useful this question is for users of this knowledge base.
Target users: {relevance_target}
Use a scale of 1 to 5, where 1 means not useful at all and 5 means extremely useful.

Provide your answer exactly in this format:

Answer:::
Evaluation: (your rationale)
Total rating: (1 to 5)

Question: {question}
Answer:::
""".strip()

QUESTION_STANDALONE_PROMPT = """
You will be given a question.
Score how understandable the question is on its own without extra context.
Use a scale of 1 to 5, where 1 means context-dependent and 5 means fully stand-alone.
If the question refers to "the context", "the document", or similar hidden context, the rating must be 1.

Provide your answer exactly in this format:

Answer:::
Evaluation: (your rationale)
Total rating: (1 to 5)

Question: {question}
Answer:::
""".strip()

QA_PATTERN = re.compile(
    r"Factoid question:\s*(.*?)\s*Answer:\s*(.*)",
    re.IGNORECASE | re.DOTALL,
)
RATING_PATTERN = re.compile(r"Total rating:\s*([1-5])", re.IGNORECASE)
EVAL_PATTERN = re.compile(
    r"Evaluation:\s*(.*?)\s*Total rating:",
    re.IGNORECASE | re.DOTALL,
)


@dataclass
class SourceDocument:
    text: str
    source: str
    title: str
    file_name: str
    source_type: str
    metadata: dict[str, Any]


@dataclass
class ChunkRecord:
    page_content: str
    metadata: dict[str, Any]


@dataclass(frozen=True)
class SubsetConfig:
    name: str
    question_type: str
    ratio: float
    max_answer_chars: int
    prompt_template: str
    context_mode: str = "single_chunk"


@dataclass(frozen=True)
class GenerationTask:
    subset: SubsetConfig
    chunks: tuple[ChunkRecord, ...]


class LLMClient:
    def generate(self, prompt: str) -> str:
        raise NotImplementedError


class HuggingFaceLLMClient(LLMClient):
    def __init__(
        self,
        *,
        model: str,
        api_key: str | None,
        timeout: float,
        max_new_tokens: int,
        temperature: float,
    ) -> None:
        self._client = InferenceClient(
            model=model,
            provider="hf-inference",
            token=api_key,
            timeout=timeout,
        )
        self._model = model
        self._max_new_tokens = max_new_tokens
        self._temperature = temperature

    def generate(self, prompt: str) -> str:
        try:
            response = self._client.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                model=self._model,
                max_tokens=self._max_new_tokens,
                temperature=self._temperature if self._temperature > 0 else None,
            )
            return str(response.choices[0].message.content or "")
        except Exception:
            response = self._client.text_generation(
                prompt,
                model=self._model,
                max_new_tokens=self._max_new_tokens,
                return_full_text=False,
                do_sample=self._temperature > 0,
                temperature=self._temperature if self._temperature > 0 else None,
            )
            return str(response or "")


class OpenAILLMClient(LLMClient):
    def __init__(
        self,
        *,
        model: str,
        api_key: str | None,
        base_url: str | None,
        timeout: float,
        max_new_tokens: int,
        temperature: float,
    ) -> None:
        self._client = OpenAI(api_key=api_key, base_url=base_url, timeout=timeout)
        self._model = model
        self._max_new_tokens = max_new_tokens
        self._temperature = temperature

    def generate(self, prompt: str) -> str:
        response = self._client.chat.completions.create(
            model=self._model,
            messages=[{"role": "user", "content": prompt}],
            temperature=self._temperature,
            max_tokens=self._max_new_tokens,
        )
        return str(response.choices[0].message.content or "")


def resolve_api_key(args: argparse.Namespace) -> str | None:
    explicit_key = str(args.api_key).strip() if args.api_key is not None else ""
    if explicit_key:
        return explicit_key

    if args.provider == "hf":
        env_key = os.environ.get("HF_TOKEN") or os.environ.get("HUGGINGFACEHUB_API_TOKEN")
        if env_key:
            return str(env_key).strip()
        raise SystemExit(
            "Missing Hugging Face token. Set HF_TOKEN in the environment or pass --api-key."
        )

    env_key = os.environ.get("OPENAI_API_KEY")
    if env_key:
        return str(env_key).strip()
    raise SystemExit(
        "Missing OpenAI-compatible API key. Set OPENAI_API_KEY in the environment or pass --api-key."
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build a synthetic RAG evaluation dataset from local documents."
    )
    source_group = parser.add_mutually_exclusive_group(required=True)
    source_group.add_argument(
        "--input-path",
        help="Source file, directory, or JSONL corpus path",
    )
    source_group.add_argument(
        "--hf-dataset",
        help="Hugging Face dataset repo id, e.g. m-ric/huggingface_doc_qa_eval",
    )
    parser.add_argument("--hf-config", default=None, help="Optional dataset config name")
    parser.add_argument("--hf-split", default="train", help="Dataset split name")
    parser.add_argument("--hf-revision", default=None, help="Optional dataset revision")
    parser.add_argument(
        "--hf-text-column",
        default=None,
        help="Column containing source text. Defaults to auto-detect.",
    )
    parser.add_argument(
        "--hf-source-column",
        default=None,
        help="Column containing source identifier/path. Defaults to auto-detect.",
    )
    parser.add_argument(
        "--hf-title-column",
        default=None,
        help="Column containing title. Defaults to auto-detect.",
    )
    parser.add_argument(
        "--hf-file-name-column",
        default=None,
        help="Column containing file name. Defaults to auto-detect.",
    )
    parser.add_argument(
        "--hf-source-type-column",
        default=None,
        help="Column containing source type. Defaults to auto-detect.",
    )
    parser.add_argument(
        "--hf-limit",
        type=int,
        default=None,
        help="Optional max number of source rows to load from the dataset",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Output directory for intermediate and final files. Defaults inside workspace-dir.",
    )
    parser.add_argument(
        "--workspace-dir",
        default=str(DEFAULT_WORKSPACE_DIR),
        help="Workspace root for Hugging Face cache and generated datasets",
    )
    parser.add_argument("--provider", choices=["hf", "openai"], default="hf")
    parser.add_argument("--model", required=True, help="LLM model id or name")
    parser.add_argument("--api-key", default=None, help="Optional API key")
    parser.add_argument("--base-url", default=None, help="Optional OpenAI-compatible base URL")
    parser.add_argument("--timeout", type=float, default=120.0)
    parser.add_argument("--temperature", type=float, default=0.0)
    parser.add_argument("--max-new-tokens", type=int, default=1000)
    parser.add_argument("--generations", type=int, default=200, help="Number of QA candidates to generate")
    parser.add_argument("--seed", type=int, default=42)
    parser.add_argument("--chunk-size", type=int, default=2000)
    parser.add_argument("--chunk-overlap", type=int, default=200)
    parser.add_argument(
        "--separators",
        nargs="*",
        default=DEFAULT_SEPARATORS[:-1],
        help=r'Separator priority, e.g. "\n\n" "\n" "." " "',
    )
    parser.add_argument("--min-context-chars", type=int, default=200)
    parser.add_argument("--max-answer-chars", type=int, default=300)
    parser.add_argument(
        "--relevance-target",
        default="users querying this knowledge base for accurate factual answers",
        help="Target audience description used by the relevance critique",
    )
    parser.add_argument("--groundedness-threshold", type=int, default=4)
    parser.add_argument("--relevance-threshold", type=int, default=4)
    parser.add_argument("--standalone-threshold", type=int, default=4)
    parser.add_argument("--unanswerable-max-groundedness", type=int, default=2)
    parser.add_argument(
        "--dataset-profile",
        choices=["legacy_single", "production_v1"],
        default="production_v1",
        help="Question-set profile for dataset generation",
    )
    parser.add_argument("--core-factoid-ratio", type=float, default=0.7)
    parser.add_argument("--cross-sentence-ratio", type=float, default=0.2)
    parser.add_argument("--unanswerable-ratio", type=float, default=0.1)
    parser.add_argument("--ambiguous-rewrite-ratio", type=float, default=0.0)
    parser.add_argument("--max-retries", type=int, default=3)
    return parser.parse_args()


def _decode_text_file(path: Path) -> str:
    payload = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "utf-16"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("latin1", errors="ignore")


def _extract_pdf_text(path: Path) -> str:
    import pymupdf

    doc = None
    texts: list[str] = []
    try:
        with suppress_stdout():
            doc = pymupdf.open(str(path))
            for page in doc:
                blocks = page.get_text("blocks")
                blocks.sort(key=lambda block: (block[1], block[0]))
                page_text = "\n".join(block[4] for block in blocks if str(block[4]).strip())
                if page_text.strip():
                    texts.append(page_text)
    finally:
        if doc is not None:
            doc.close()
    return "\n\n".join(texts)


def _load_jsonl_corpus(path: Path) -> list[SourceDocument]:
    docs: list[SourceDocument] = []
    with path.open("r", encoding="utf-8") as handle:
        for idx, line in enumerate(handle):
            line = line.strip()
            if not line:
                continue
            row = json.loads(line)
            text = (
                row.get("contents")
                or row.get("text")
                or row.get("page_content")
                or row.get("content")
                or ""
            )
            text = reflow_paragraphs(clean_text(str(text or "")))
            if not text:
                continue
            source = str(row.get("source") or row.get("source_uri") or path.resolve())
            title = str(row.get("title") or row.get("file_name") or path.stem or f"row-{idx}")
            file_name = str(row.get("file_name") or Path(source).name or path.name)
            source_type = str(row.get("source_type") or path.suffix.lstrip(".") or "jsonl")
            metadata = {k: v for k, v in row.items() if k not in {"contents", "text", "page_content", "content"}}
            docs.append(
                SourceDocument(
                    text=text,
                    source=source,
                    title=title,
                    file_name=file_name,
                    source_type=source_type,
                    metadata=metadata,
                )
            )
    return docs


def _load_excel_documents(path: Path) -> list[SourceDocument]:
    excel_data = pd.read_excel(path, sheet_name=None, dtype=str)
    docs: list[SourceDocument] = []
    for sheet_name, frame in excel_data.items():
        frame = frame.fillna("")
        columns = frame.columns.tolist()
        for idx, row in frame.iterrows():
            content_parts: list[str] = []
            for column in columns:
                value = str(row[column]).strip()
                if value:
                    content_parts.append(f"{column}={value}")
            if not content_parts:
                continue
            docs.append(
                SourceDocument(
                    text="；".join(content_parts),
                    source=str(path.resolve()),
                    title=f"{path.stem} / {sheet_name}",
                    file_name=path.name,
                    source_type="excel",
                    metadata={"sheet_name": sheet_name, "row_index": int(idx)},
                )
            )
    return docs


def _load_source_document(path: Path) -> list[SourceDocument]:
    suffix = path.suffix.lower()
    if suffix in JSONL_EXTS:
        return _load_jsonl_corpus(path)
    if suffix in EXCEL_EXTS:
        return _load_excel_documents(path)
    if suffix in TEXT_EXTS:
        text = _decode_text_file(path)
    elif suffix in DOCX_EXTS:
        text = _read_docx_text(str(path)) or ""
    elif suffix in WORD_LEGACY_EXTS:
        text = _read_via_office_convert(str(path)) or ""
    elif suffix in PDFLIKE_EXTS:
        text = _extract_pdf_text(path)
    else:
        return []

    normalized = reflow_paragraphs(clean_text(text))
    if not normalized:
        return []
    return [
        SourceDocument(
            text=normalized,
            source=str(path.resolve()),
            title=path.stem,
            file_name=path.name,
            source_type=suffix.lstrip("."),
            metadata={},
        )
    ]


def load_source_documents(input_path: Path) -> list[SourceDocument]:
    if input_path.is_file():
        return _load_source_document(input_path)

    docs: list[SourceDocument] = []
    for file_path in sorted(p for p in input_path.rglob("*") if p.is_file()):
        docs.extend(_load_source_document(file_path))
    return docs


def build_subset_configs(args: argparse.Namespace) -> list[SubsetConfig]:
    if args.dataset_profile == "legacy_single":
        return [
            SubsetConfig(
                name="legacy_single",
                question_type="answerable",
                ratio=1.0,
                max_answer_chars=args.max_answer_chars,
                prompt_template=QA_GENERATION_PROMPT,
            )
        ]

    configs = [
        SubsetConfig(
            name="core_factoid",
            question_type="answerable",
            ratio=args.core_factoid_ratio,
            max_answer_chars=250,
            prompt_template=CORE_FACTOID_PROMPT,
        ),
        SubsetConfig(
            name="cross_sentence",
            question_type="answerable",
            ratio=args.cross_sentence_ratio,
            max_answer_chars=300,
            prompt_template=CROSS_SENTENCE_PROMPT,
            context_mode="adjacent_chunk_pair",
        ),
        SubsetConfig(
            name="unanswerable",
            question_type="unanswerable",
            ratio=args.unanswerable_ratio,
            max_answer_chars=250,
            prompt_template=UNANSWERABLE_PROMPT,
        ),
        SubsetConfig(
            name="ambiguous_rewrite",
            question_type="answerable",
            ratio=args.ambiguous_rewrite_ratio,
            max_answer_chars=250,
            prompt_template=AMBIGUOUS_REWRITE_PROMPT,
        ),
    ]
    active = [config for config in configs if config.ratio > 0]
    if not active:
        raise SystemExit("At least one active subset ratio is required.")
    return active


def slugify_name(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9._-]+", "_", str(value or "").strip())
    slug = slug.strip("._-")
    return slug or "dataset"


def configure_workspace(workspace_dir: Path) -> dict[str, Path]:
    hf_home = workspace_dir / "hf_cache"
    hf_hub_cache = hf_home / "hub"
    hf_datasets_cache = hf_home / "datasets"
    prepared_root = workspace_dir / "prepared"
    output_root = workspace_dir / "generated"
    for path in (workspace_dir, hf_home, hf_hub_cache, hf_datasets_cache, prepared_root, output_root):
        path.mkdir(parents=True, exist_ok=True)

    os.environ["HF_HOME"] = str(hf_home)
    os.environ["HF_HUB_CACHE"] = str(hf_hub_cache)
    os.environ["HF_DATASETS_CACHE"] = str(hf_datasets_cache)
    return {
        "workspace_dir": workspace_dir,
        "hf_home": hf_home,
        "hf_hub_cache": hf_hub_cache,
        "hf_datasets_cache": hf_datasets_cache,
        "prepared_root": prepared_root,
        "output_root": output_root,
    }


def _pick_column(columns: Sequence[str], preferred: str | None, candidates: Sequence[str]) -> str | None:
    if preferred:
        if preferred not in columns:
            raise SystemExit(f"Column '{preferred}' not found. Available columns: {', '.join(columns)}")
        return preferred
    for candidate in candidates:
        if candidate in columns:
            return candidate
    return None


def load_hf_source_documents(args: argparse.Namespace) -> list[SourceDocument]:
    from datasets import load_dataset

    dataset = load_dataset(
        path=args.hf_dataset,
        name=args.hf_config,
        split=args.hf_split,
        revision=args.hf_revision,
        cache_dir=str(args.hf_datasets_cache_dir),
    )
    if args.hf_limit is not None:
        dataset = dataset.select(range(min(args.hf_limit, len(dataset))))

    columns = list(dataset.column_names)
    text_column = _pick_column(columns, args.hf_text_column, DEFAULT_TEXT_COLUMNS)
    if text_column is None:
        raise SystemExit(
            "Unable to infer dataset text column. "
            f"Pass --hf-text-column explicitly. Available columns: {', '.join(columns)}"
        )
    source_column = _pick_column(columns, args.hf_source_column, DEFAULT_SOURCE_COLUMNS)
    title_column = _pick_column(columns, args.hf_title_column, DEFAULT_TITLE_COLUMNS)
    file_name_column = _pick_column(columns, args.hf_file_name_column, DEFAULT_FILE_NAME_COLUMNS)
    source_type_column = _pick_column(
        columns,
        args.hf_source_type_column,
        DEFAULT_SOURCE_TYPE_COLUMNS,
    )

    docs: list[SourceDocument] = []
    excluded_columns = {
        text_column,
        source_column,
        title_column,
        file_name_column,
        source_type_column,
    }
    for idx, row in enumerate(dataset):
        text = reflow_paragraphs(clean_text(str(row.get(text_column) or "")))
        if not text:
            continue
        source = str(
            row.get(source_column)
            or f"hf://{args.hf_dataset}/{args.hf_split}/{idx}"
        )
        title = str(
            row.get(title_column)
            or Path(source).stem
            or f"{args.hf_dataset}-{idx}"
        )
        file_name = str(
            row.get(file_name_column)
            or Path(source).name
            or f"{args.hf_dataset.replace('/', '_')}-{idx}.txt"
        )
        source_type = str(row.get(source_type_column) or "hf_dataset")
        metadata = {
            k: v
            for k, v in row.items()
            if k not in excluded_columns and v not in (None, "")
        }
        metadata.update(
            {
                "hf_dataset": args.hf_dataset,
                "hf_config": args.hf_config,
                "hf_split": args.hf_split,
                "hf_row_index": idx,
            }
        )
        docs.append(
            SourceDocument(
                text=text,
                source=source,
                title=title,
                file_name=file_name,
                source_type=source_type,
                metadata=metadata,
            )
        )
    return docs


def load_source_documents_from_jsonl(path: Path) -> list[SourceDocument]:
    docs: list[SourceDocument] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            row = json.loads(line)
            docs.append(
                SourceDocument(
                    text=str(row.get("text") or ""),
                    source=str(row.get("source") or ""),
                    title=str(row.get("title") or ""),
                    file_name=str(row.get("file_name") or ""),
                    source_type=str(row.get("source_type") or ""),
                    metadata=dict(row.get("metadata") or {}),
                )
            )
    return docs


def load_chunks_from_jsonl(path: Path) -> list[ChunkRecord]:
    chunks: list[ChunkRecord] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            row = json.loads(line)
            chunks.append(
                ChunkRecord(
                    page_content=str(row.get("page_content") or ""),
                    metadata=dict(row.get("metadata") or {}),
                )
            )
    return chunks


def save_json(obj: Any, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        json.dumps(obj, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def stable_payload_hash(payload: Any) -> str:
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:12]


def build_source_cache_manifest(
    args: argparse.Namespace,
    *,
    input_path: Path | None,
    source_label: str,
) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "source_label": source_label,
        "hf_dataset": args.hf_dataset,
        "hf_config": args.hf_config,
        "hf_split": args.hf_split if args.hf_dataset else None,
        "hf_revision": args.hf_revision,
        "hf_limit": args.hf_limit,
        "hf_text_column": args.hf_text_column,
        "hf_source_column": args.hf_source_column,
        "hf_title_column": args.hf_title_column,
        "hf_file_name_column": args.hf_file_name_column,
        "hf_source_type_column": args.hf_source_type_column,
        "input_path": str(input_path) if input_path else None,
    }


def build_chunk_cache_config(args: argparse.Namespace) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "chunk_pool_mode": "shared",
        "chunk_size": args.chunk_size,
        "chunk_overlap": args.chunk_overlap,
        "min_context_chars": args.min_context_chars,
        "separators": list(args.separators) + [""],
    }


def build_generation_run_config(
    args: argparse.Namespace,
    subsets: Sequence[SubsetConfig],
) -> dict[str, Any]:
    return {
        "schema_version": 1,
        "dataset_profile": args.dataset_profile,
        "provider": args.provider,
        "model": args.model,
        "max_new_tokens": args.max_new_tokens,
        "temperature": args.temperature,
        "generations": args.generations,
        "seed": args.seed,
        "subset_configs": subset_config_summary(subsets),
        "thresholds": {
            "groundedness": args.groundedness_threshold,
            "relevance": args.relevance_threshold,
            "standalone": args.standalone_threshold,
            "unanswerable_max_groundedness": args.unanswerable_max_groundedness,
        },
    }


def build_generation_run_slug(
    args: argparse.Namespace,
    subsets: Sequence[SubsetConfig],
) -> str:
    run_config = build_generation_run_config(args, subsets)
    token = stable_payload_hash(run_config)
    model_slug = slugify_name(args.model)
    return f"{args.dataset_profile}__g{args.generations}__{model_slug}__{token}"


def ensure_prepared_source_documents(
    *,
    args: argparse.Namespace,
    workspace_paths: dict[str, Path],
    dataset_slug: str,
    input_path: Path | None,
    source_label: str,
) -> tuple[list[SourceDocument], Path, dict[str, Any], bool]:
    prepared_dir = workspace_paths["prepared_root"] / dataset_slug
    prepared_dir.mkdir(parents=True, exist_ok=True)
    source_docs_path = prepared_dir / "source_documents.jsonl"
    source_manifest_path = prepared_dir / "source_documents.manifest.json"
    expected_manifest = build_source_cache_manifest(args, input_path=input_path, source_label=source_label)

    if source_docs_path.exists() and source_manifest_path.exists():
        actual_manifest = json.loads(source_manifest_path.read_text(encoding="utf-8"))
        if actual_manifest == expected_manifest:
            LOGGER.info("using prepared source documents cache %s", source_docs_path)
            return load_source_documents_from_jsonl(source_docs_path), source_docs_path, expected_manifest, True

    if args.hf_dataset:
        LOGGER.info("loading source documents from dataset %s", source_label)
        source_docs = load_hf_source_documents(args)
    else:
        assert input_path is not None
        LOGGER.info("loading source documents from %s", input_path)
        source_docs = load_source_documents(input_path)

    save_jsonl((asdict(doc) for doc in source_docs), source_docs_path)
    save_json(expected_manifest, source_manifest_path)
    return source_docs, source_docs_path, expected_manifest, False


def ensure_prepared_chunk_pool(
    *,
    args: argparse.Namespace,
    workspace_paths: dict[str, Path],
    dataset_slug: str,
    source_docs: Sequence[SourceDocument],
    source_manifest: dict[str, Any],
) -> tuple[list[ChunkRecord], Path, dict[str, Any], bool]:
    prepared_dir = workspace_paths["prepared_root"] / dataset_slug
    chunk_config = build_chunk_cache_config(args)
    chunk_token = stable_payload_hash(chunk_config)
    chunk_dir = prepared_dir / "chunks" / f"shared__c{args.chunk_size}__o{args.chunk_overlap}__m{args.min_context_chars}__{chunk_token}"
    chunk_dir.mkdir(parents=True, exist_ok=True)
    contexts_path = chunk_dir / "contexts.jsonl"
    chunk_manifest_path = chunk_dir / "manifest.json"
    expected_manifest = {
        "schema_version": 1,
        "source_manifest": source_manifest,
        "chunk_config": chunk_config,
    }

    if contexts_path.exists() and chunk_manifest_path.exists():
        actual_manifest = json.loads(chunk_manifest_path.read_text(encoding="utf-8"))
        if actual_manifest == expected_manifest:
            LOGGER.info("using prepared chunk pool cache %s", contexts_path)
            return load_chunks_from_jsonl(contexts_path), contexts_path, expected_manifest, True

    chunks = build_shared_chunk_pool(
        source_docs,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        separators=chunk_config["separators"],
        min_context_chars=args.min_context_chars,
    )
    chunk_rows = [
        {"page_content": chunk.page_content, "metadata": dict(chunk.metadata)}
        for chunk in chunks
    ]
    save_jsonl(chunk_rows, contexts_path)
    save_json(expected_manifest, chunk_manifest_path)
    return chunks, contexts_path, expected_manifest, False


def split_document_text(
    text: str,
    *,
    chunk_size: int,
    chunk_overlap: int,
    separators: Sequence[str],
) -> list[tuple[str, int]]:
    text = str(text or "")
    if not text.strip():
        return []

    normalized_separators = list(separators) + [""]
    chunks: list[tuple[str, int]] = []
    start = 0
    text_length = len(text)
    half_window = max(1, int(chunk_size * 0.5))

    while start < text_length:
        max_end = min(start + chunk_size, text_length)
        end = max_end
        if max_end < text_length:
            for separator in normalized_separators:
                if not separator:
                    continue
                candidate = text.rfind(separator, start + half_window, max_end)
                if candidate != -1:
                    end = candidate + len(separator)
                    break
        if end <= start:
            end = max_end

        raw_slice = text[start:end]
        stripped = raw_slice.strip()
        if stripped:
            left_trim = len(raw_slice) - len(raw_slice.lstrip())
            actual_start = start + left_trim
            chunks.append((stripped, actual_start))

        if end >= text_length:
            break
        next_start = max(start + 1, end - chunk_overlap)
        if next_start <= start:
            next_start = end
        start = next_start
    return chunks


def build_chunks(
    source_docs: Sequence[SourceDocument],
    *,
    chunk_size: int,
    chunk_overlap: int,
    separators: Sequence[str],
    min_context_chars: int,
    chunk_profile_name: str | None = None,
) -> list[ChunkRecord]:
    processed: list[ChunkRecord] = []
    seen_keys: set[tuple[str, int, str]] = set()
    for doc in source_docs:
        for chunk_text, start_index in split_document_text(
            doc.text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
        ):
            if len(chunk_text) < min_context_chars:
                continue
            dedupe_key = (doc.source, start_index, chunk_text)
            if dedupe_key in seen_keys:
                continue
            seen_keys.add(dedupe_key)
            metadata = dict(doc.metadata)
            metadata.update(
                {
                    "source": doc.source,
                    "title": doc.title,
                    "file_name": doc.file_name,
                    "source_type": doc.source_type,
                    "start_index": start_index,
                }
            )
            if chunk_profile_name:
                metadata["chunk_profile"] = chunk_profile_name
            processed.append(ChunkRecord(page_content=chunk_text, metadata=metadata))
    return processed


def build_llm_client(args: argparse.Namespace) -> LLMClient:
    api_key = resolve_api_key(args)
    if args.provider == "hf":
        return HuggingFaceLLMClient(
            model=args.model,
            api_key=api_key,
            timeout=args.timeout,
            max_new_tokens=args.max_new_tokens,
            temperature=args.temperature,
        )
    return OpenAILLMClient(
        model=args.model,
        api_key=api_key,
        base_url=args.base_url,
        timeout=args.timeout,
        max_new_tokens=args.max_new_tokens,
        temperature=args.temperature,
    )


def call_llm_with_retry(
    client: LLMClient,
    prompt: str,
    *,
    max_retries: int,
) -> str:
    last_error: Exception | None = None
    for attempt in range(1, max_retries + 1):
        try:
            return client.generate(prompt)
        except Exception as exc:  # noqa: BLE001
            last_error = exc
            if attempt >= max_retries:
                break
            sleep_seconds = min(8.0, float(attempt))
            LOGGER.warning("llm call failed attempt=%s error=%s", attempt, exc)
            time.sleep(sleep_seconds)
    assert last_error is not None
    raise last_error


def parse_qa_output(text: str) -> tuple[str, str] | None:
    matches = list(QA_PATTERN.finditer(str(text or "")))
    if not matches:
        return None
    match = matches[-1]
    question = match.group(1).strip()
    answer = match.group(2).strip()
    if not question or not answer:
        return None
    return question, answer


def parse_rating_output(text: str) -> tuple[int, str] | None:
    rating_match = RATING_PATTERN.search(str(text or ""))
    if not rating_match:
        return None
    evaluation_match = EVAL_PATTERN.search(str(text or ""))
    evaluation = evaluation_match.group(1).strip() if evaluation_match else ""
    return int(rating_match.group(1)), evaluation


def sample_items(
    items: Sequence[Any],
    *,
    n_generations: int,
    rng: random.Random,
) -> list[Any]:
    if not items:
        return []
    if n_generations <= len(items):
        return rng.sample(list(items), n_generations)
    return [rng.choice(list(items)) for _ in range(n_generations)]


def allocate_subset_generations(
    subsets: Sequence[SubsetConfig],
    total_generations: int,
) -> dict[str, int]:
    if total_generations <= 0:
        return {subset.name: 0 for subset in subsets}
    total_ratio = sum(max(0.0, subset.ratio) for subset in subsets)
    if total_ratio <= 0:
        raise SystemExit("Subset ratios must sum to a positive value.")

    raw_allocations: list[tuple[str, int, float]] = []
    assigned = 0
    for subset in subsets:
        exact = total_generations * (subset.ratio / total_ratio)
        base = int(exact)
        raw_allocations.append((subset.name, base, exact - base))
        assigned += base

    remaining = total_generations - assigned
    allocations = {name: base for name, base, _ in raw_allocations}
    for name, _, _ in sorted(raw_allocations, key=lambda item: item[2], reverse=True)[:remaining]:
        allocations[name] += 1
    return allocations


def build_shared_chunk_pool(
    source_docs: Sequence[SourceDocument],
    *,
    chunk_size: int,
    chunk_overlap: int,
    separators: Sequence[str],
    min_context_chars: int,
) -> list[ChunkRecord]:
    return build_chunks(
        source_docs,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
        min_context_chars=min_context_chars,
        chunk_profile_name="shared",
    )


def build_adjacent_chunk_pairs(chunk_pool: Sequence[ChunkRecord]) -> list[tuple[ChunkRecord, ChunkRecord]]:
    chunks_by_source: dict[str, list[ChunkRecord]] = defaultdict(list)
    for chunk in chunk_pool:
        source = str(chunk.metadata.get("source") or "")
        if source:
            chunks_by_source[source].append(chunk)

    pairs: list[tuple[ChunkRecord, ChunkRecord]] = []
    for source_chunks in chunks_by_source.values():
        ordered = sorted(source_chunks, key=lambda chunk: int(chunk.metadata.get("start_index") or 0))
        for left, right in zip(ordered, ordered[1:]):
            left_start = int(left.metadata.get("start_index") or 0)
            right_start = int(right.metadata.get("start_index") or 0)
            if left_start == right_start:
                continue
            pairs.append((left, right))
    return pairs


def build_generation_tasks(
    *,
    chunk_pool: Sequence[ChunkRecord],
    subsets: Sequence[SubsetConfig],
    total_generations: int,
    rng: random.Random,
) -> list[GenerationTask]:
    allocations = allocate_subset_generations(subsets, total_generations)
    tasks: list[GenerationTask] = []
    adjacent_pairs = build_adjacent_chunk_pairs(chunk_pool)
    for subset in subsets:
        if subset.context_mode == "adjacent_chunk_pair":
            if allocations.get(subset.name, 0) > 0 and not adjacent_pairs:
                LOGGER.warning("no adjacent chunk pairs available for subset=%s", subset.name)
            sampled = sample_items(
                adjacent_pairs,
                n_generations=allocations.get(subset.name, 0),
                rng=rng,
            )
            for pair in sampled:
                tasks.append(GenerationTask(subset=subset, chunks=pair))
            continue

        sampled = sample_items(
            chunk_pool,
            n_generations=allocations.get(subset.name, 0),
            rng=rng,
        )
        for chunk in sampled:
            tasks.append(GenerationTask(subset=subset, chunks=(chunk,)))
    rng.shuffle(tasks)
    return tasks


def format_task_context(chunks: Sequence[ChunkRecord]) -> str:
    if len(chunks) == 1:
        return chunks[0].page_content
    return "\n\n".join(
        f"Context {idx}:\n{chunk.page_content}"
        for idx, chunk in enumerate(chunks, start=1)
    )


def generate_candidates(
    tasks: Sequence[GenerationTask],
    *,
    client: LLMClient,
    args: argparse.Namespace,
) -> list[dict[str, Any]]:
    outputs: list[dict[str, Any]] = []
    for index, task in enumerate(tasks, start=1):
        LOGGER.info("generate qa %s/%s", index, len(tasks))
        subset = task.subset
        chunks = task.chunks
        primary_chunk = chunks[0]
        context_text = format_task_context(chunks)
        raw_output = call_llm_with_retry(
            client,
            subset.prompt_template.format(context=context_text),
            max_retries=args.max_retries,
        )
        parsed = parse_qa_output(raw_output)
        if parsed is None:
            continue
        question, answer = parsed
        if subset.question_type == "unanswerable":
            answer = "NO_ANSWER"
        if len(answer) > subset.max_answer_chars:
            continue
        question_id = stable_payload_hash(
            {
                "subset_name": subset.name,
                "question": question,
                "answer": answer,
                "context_sources": [chunk.metadata.get("source") for chunk in chunks],
                "context_chunk_start_indexes": [
                    chunk.metadata.get("start_index") for chunk in chunks
                ],
            }
        )
        outputs.append(
            {
                "question_id": question_id,
                "qid": question_id,
                "context": context_text,
                "question": question,
                "answer": answer,
                "golden_answers": [answer],
                "subset_name": subset.name,
                "question_type": subset.question_type,
                "context_mode": subset.context_mode,
                "context_count": len(chunks),
                "chunk_profile": primary_chunk.metadata.get("chunk_profile"),
                "source_doc": primary_chunk.metadata.get("source"),
                "file_name": primary_chunk.metadata.get("file_name"),
                "title": primary_chunk.metadata.get("title"),
                "source_type": primary_chunk.metadata.get("source_type"),
                "chunk_start_index": primary_chunk.metadata.get("start_index"),
                "chunk_metadata": dict(primary_chunk.metadata),
                "context_sources": [chunk.metadata.get("source") for chunk in chunks],
                "context_chunk_start_indexes": [chunk.metadata.get("start_index") for chunk in chunks],
                "context_chunks": [
                    {
                        "page_content": chunk.page_content,
                        "metadata": dict(chunk.metadata),
                    }
                    for chunk in chunks
                ],
                "generation_raw_output": raw_output,
            }
        )
    return outputs


def critique_candidates(
    candidates: list[dict[str, Any]],
    *,
    client: LLMClient,
    args: argparse.Namespace,
) -> None:
    for index, item in enumerate(candidates, start=1):
        LOGGER.info("critique qa %s/%s", index, len(candidates))
        prompts = {
            "groundedness": QUESTION_GROUNDEDNESS_PROMPT.format(
                context=item["context"],
                question=item["question"],
            ),
            "relevance": QUESTION_RELEVANCE_PROMPT.format(
                question=item["question"],
                relevance_target=args.relevance_target,
            ),
            "standalone": QUESTION_STANDALONE_PROMPT.format(question=item["question"]),
        }
        for criterion, prompt in prompts.items():
            raw_output = call_llm_with_retry(client, prompt, max_retries=args.max_retries)
            parsed = parse_rating_output(raw_output)
            item[f"{criterion}_raw_output"] = raw_output
            if parsed is None:
                continue
            score, evaluation = parsed
            item[f"{criterion}_score"] = score
            item[f"{criterion}_eval"] = evaluation


def filter_candidates(candidates: Sequence[dict[str, Any]], args: argparse.Namespace) -> list[dict[str, Any]]:
    filtered: list[dict[str, Any]] = []
    for item in candidates:
        groundedness = int(item.get("groundedness_score") or 0)
        relevance = int(item.get("relevance_score") or 0)
        standalone = int(item.get("standalone_score") or 0)
        question_type = str(item.get("question_type") or "answerable")
        if question_type == "unanswerable":
            if groundedness > args.unanswerable_max_groundedness:
                continue
            if str(item.get("answer") or "").strip() != "NO_ANSWER":
                continue
        elif groundedness < args.groundedness_threshold:
            continue
        if relevance < args.relevance_threshold:
            continue
        if standalone < args.standalone_threshold:
            continue
        filtered.append(item)
    return filtered


def count_by_subset(rows: Sequence[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for row in rows:
        subset_name = str(row.get("subset_name") or "unknown")
        counts[subset_name] = counts.get(subset_name, 0) + 1
    return counts


def subset_config_summary(subsets: Sequence[SubsetConfig]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for subset in subsets:
        rows.append(
            {
                "name": subset.name,
                "question_type": subset.question_type,
                "ratio": subset.ratio,
                "max_answer_chars": subset.max_answer_chars,
                "context_mode": subset.context_mode,
            }
        )
    return rows


def save_jsonl(rows: Iterable[dict[str, Any]], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8", newline="\n") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def save_csv(rows: Sequence[dict[str, Any]], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pd.DataFrame.from_records(list(rows)).to_csv(output_path, index=False)


def main() -> None:
    args = parse_args()
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    )

    rng = random.Random(args.seed)
    workspace_dir = Path(args.workspace_dir).resolve()
    workspace_paths = configure_workspace(workspace_dir)
    args.hf_datasets_cache_dir = workspace_paths["hf_datasets_cache"]
    input_path: Path | None = None
    source_label: str
    dataset_slug: str
    if args.hf_dataset:
        dataset_slug = slugify_name(args.hf_dataset)
        source_label = f"hf://{args.hf_dataset}/{args.hf_split}"
    else:
        input_path = Path(args.input_path).resolve()
        dataset_slug = slugify_name(input_path.name)
        source_label = str(input_path)

    source_docs, prepared_source_docs_path, source_manifest, source_cache_hit = ensure_prepared_source_documents(
        args=args,
        workspace_paths=workspace_paths,
        dataset_slug=dataset_slug,
        input_path=input_path,
        source_label=source_label,
    )
    if not source_docs:
        raise SystemExit(f"No supported documents loaded from {source_label}")

    subset_configs = build_subset_configs(args)
    output_dir = (
        Path(args.output_dir).resolve()
        if args.output_dir
        else workspace_paths["output_root"] / dataset_slug / build_generation_run_slug(args, subset_configs)
    )
    output_dir.mkdir(parents=True, exist_ok=True)

    chunks, prepared_contexts_path, chunk_manifest, chunk_cache_hit = ensure_prepared_chunk_pool(
        args=args,
        workspace_paths=workspace_paths,
        dataset_slug=dataset_slug,
        source_docs=source_docs,
        source_manifest=source_manifest,
    )
    if not chunks:
        raise SystemExit("No chunks generated. Adjust chunk or min-context settings.")

    adjacent_chunk_pair_count = len(build_adjacent_chunk_pairs(chunks))
    generation_tasks = build_generation_tasks(
        chunk_pool=chunks,
        subsets=subset_configs,
        total_generations=args.generations,
        rng=rng,
    )
    client = build_llm_client(args)
    candidates = generate_candidates(generation_tasks, client=client, args=args)
    if not candidates:
        raise SystemExit("No QA candidates generated.")

    critique_candidates(candidates, client=client, args=args)
    filtered = filter_candidates(candidates, args)

    save_jsonl(candidates, output_dir / "qa_candidates.jsonl")
    save_csv(candidates, output_dir / "qa_candidates.csv")
    save_jsonl(filtered, output_dir / "qa_eval_dataset.jsonl")
    save_csv(filtered, output_dir / "qa_eval_dataset.csv")

    summary = {
        "input_path": str(input_path) if input_path else None,
        "hf_dataset": args.hf_dataset,
        "hf_config": args.hf_config,
        "hf_split": args.hf_split if args.hf_dataset else None,
        "source_label": source_label,
        "workspace_dir": str(workspace_dir),
        "hf_home": str(workspace_paths["hf_home"]),
        "hf_hub_cache": str(workspace_paths["hf_hub_cache"]),
        "hf_datasets_cache": str(workspace_paths["hf_datasets_cache"]),
        "prepared_dir": str(workspace_paths["prepared_root"] / dataset_slug),
        "prepared_source_documents_path": str(prepared_source_docs_path),
        "prepared_contexts_path": str(prepared_contexts_path),
        "source_cache_hit": source_cache_hit,
        "chunk_cache_hit": chunk_cache_hit,
        "output_dir": str(output_dir),
        "dataset_profile": args.dataset_profile,
        "source_document_count": len(source_docs),
        "context_count": len(chunks),
        "sampled_task_count": len(generation_tasks),
        "sampled_context_count": len(generation_tasks),
        "candidate_count": len(candidates),
        "filtered_count": len(filtered),
        "provider": args.provider,
        "model": args.model,
        "chunk_size": args.chunk_size,
        "chunk_overlap": args.chunk_overlap,
        "subset_configs": subset_config_summary(subset_configs),
        "source_manifest": source_manifest,
        "chunk_manifest": chunk_manifest,
        "chunk_pool_mode": "shared",
        "shared_chunk_count": len(chunks),
        "adjacent_chunk_pair_count": adjacent_chunk_pair_count,
        "subset_candidate_counts": count_by_subset(candidates),
        "subset_filtered_counts": count_by_subset(filtered),
        "thresholds": {
            "groundedness": args.groundedness_threshold,
            "relevance": args.relevance_threshold,
            "standalone": args.standalone_threshold,
            "unanswerable_max_groundedness": args.unanswerable_max_groundedness,
        },
    }
    save_jsonl([summary], output_dir / "summary.jsonl")
    (output_dir / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    LOGGER.info("done source_docs=%s contexts=%s candidates=%s filtered=%s", len(source_docs), len(chunks), len(candidates), len(filtered))
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
